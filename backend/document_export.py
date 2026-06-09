"""Markdown → Word (.docx) export helpers."""
from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path


def is_docx_export_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        pass
    return bool(shutil.which("pandoc"))


def markdown_to_docx_bytes(markdown: str, *, title: str | None = None) -> bytes:
    md = (markdown or "").strip()
    if not md:
        raise ValueError("文稿内容为空，无法导出 Word")
    if shutil.which("pandoc"):
        try:
            return _markdown_to_docx_pandoc(md, title=title)
        except Exception:
            pass
    return _markdown_to_docx_python(md, title=title)


def _markdown_to_docx_pandoc(md: str, *, title: str | None = None) -> bytes:
    body = f"# {title}\n\n{md}" if title else md
    with tempfile.TemporaryDirectory() as td:
        md_path = Path(td) / "input.md"
        out_path = Path(td) / "output.docx"
        md_path.write_text(body, encoding="utf-8")
        proc = subprocess.run(
            ["pandoc", str(md_path), "-o", str(out_path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or "pandoc 转换失败").strip()
            raise RuntimeError(err)
        return out_path.read_bytes()


def _add_md_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _markdown_to_docx_python(md: str, *, title: str | None = None) -> bytes:
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    blocks = re.split(r"\n\n+", md)
    for block in blocks:
        line = block.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line, re.DOTALL)
        if heading:
            level = min(len(heading.group(1)), 9)
            doc.add_heading(heading.group(2).strip(), level=level)
            continue
        for para in line.split("\n"):
            p = para.strip()
            if not p:
                continue
            if p.startswith("---") or p.startswith("***"):
                continue
            paragraph = doc.add_paragraph()
            _add_md_runs(paragraph, p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_download_stem(name: str, *, fallback: str = "export") -> str:
    stem = Path(name).stem if name else fallback
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", stem).strip("._") or fallback
    return safe[:80]
