"""导入/上传多格式集成测试（需本地后端运行）。"""
from __future__ import annotations

import io
import uuid
import zipfile

import pytest

from helpers.client import ApiClient


def _create_upload_project(client: ApiClient) -> dict:
    title = f"pytest-upload-{uuid.uuid4().hex[:8]}"
    resp = client.post("/projects", json_body={"title": title, "genre": "测试", "work_type": "novel_long"})
    assert resp.status_code == 200, resp.text[:300]
    return resp.json()


def _upload(client: ApiClient, project_id: str, name: str, data: bytes, subdir: str = "正文"):
    return client.post(
        f"/projects/{project_id}/upload",
        files={"file": (name, data, "application/octet-stream")},
        query={"subdir": subdir},
        timeout=120,
    )


def _zip_bytes(entries: dict[str, bytes | str]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for path, content in entries.items():
            payload = content.encode("utf-8") if isinstance(content, str) else content
            zf.writestr(path, payload)
    return buf.getvalue()


@pytest.fixture
def upload_project(admin_client):
    project = _create_upload_project(admin_client)
    yield project
    admin_client.delete(f"/projects/{project['id']}")


class TestUploadTextFormats:
    def test_upload_md_to_manuscript(self, admin_client, upload_project):
        pid = upload_project["id"]
        body = "# 第一章\n\n导入测试正文。".encode("utf-8")
        resp = _upload(admin_client, pid, "chapter-one.md", body, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert data.get("status") == "ok"
        assert any(c["filename"].endswith(".md") for c in data.get("chapters", []))

    def test_upload_txt_to_manuscript(self, admin_client, upload_project):
        pid = upload_project["id"]
        body = "第二章\n\n纯文本导入测试。".encode("gbk")
        resp = _upload(admin_client, pid, "chapter-two.txt", body, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert data.get("status") == "ok"
        assert len(data.get("chapters", [])) >= 1

    def test_upload_md_to_outline(self, admin_client, upload_project):
        pid = upload_project["id"]
        body = "# 故事大纲\n\n起承转合。".encode("utf-8")
        resp = _upload(admin_client, pid, "outline.md", body, "大纲")
        assert resp.status_code == 200, resp.text[:400]
        files = admin_client.get(f"/projects/{pid}/files/outline").json()
        assert any(f["filename"].endswith(".md") for f in files)


class TestUploadZipFormats:
    def test_zip_standard_body_layout(self, admin_client, upload_project):
        pid = upload_project["id"]
        z = _zip_bytes({"正文/zip-ch1.md": "# Zip章一\n\n内容。"})
        resp = _upload(admin_client, pid, "standard.zip", z, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) >= 1
        assert not data.get("import_warning")

    def test_zip_wrapped_folder(self, admin_client, upload_project):
        pid = upload_project["id"]
        z = _zip_bytes({"my-book/正文/zip-ch2.md": "# Zip章二\n\n内容。"})
        resp = _upload(admin_client, pid, "wrapped.zip", z, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) >= 1, data

    def test_zip_root_markdown_hoisted(self, admin_client, upload_project):
        pid = upload_project["id"]
        z = _zip_bytes({"root-ch.md": "# 根目录章\n\n应归入正文。"})
        resp = _upload(admin_client, pid, "root-md.zip", z, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) >= 1, data

    def test_zip_no_manuscript_warns(self, admin_client, upload_project):
        pid = upload_project["id"]
        z = _zip_bytes({"assets/cover.bin": b"\x00\x01binary"})
        resp = _upload(admin_client, pid, "no-chapter.zip", z, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) == 0
        assert data.get("import_warning")


class TestUploadPythonFormats:
    @pytest.fixture(scope="class")
    def doc_libs(self):
        try:
            import mammoth  # noqa: F401
            import fitz  # noqa: F401
            import markdownify  # noqa: F401
        except ImportError:
            pytest.skip("未安装 docx/pdf/html 转换依赖：pip install -r backend/requirements.txt")
        return True

    def test_upload_html(self, admin_client, upload_project, doc_libs):
        pid = upload_project["id"]
        html = b"<html><body><h1>HTML \u7ae0</h1><p>\u6bb5\u843d\u3002</p></body></html>"
        resp = _upload(admin_client, pid, "chapter.html", html, "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) >= 1

    def test_upload_docx(self, admin_client, upload_project, doc_libs):
        from docx import Document

        pid = upload_project["id"]
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading("Docx \u7ae0", level=1)
        doc.add_paragraph("\u8fd9\u662f docx \u5bfc\u5165\u6d4b\u8bd5\u3002")
        doc.save(buf)
        resp = _upload(admin_client, pid, "chapter.docx", buf.getvalue(), "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) >= 1

    def test_upload_pdf(self, admin_client, upload_project, doc_libs):
        import fitz

        pid = upload_project["id"]
        buf = io.BytesIO()
        pdf = fitz.open()
        page = pdf.new_page()
        page.insert_text((72, 72), "PDF import test chapter.")
        pdf.save(buf)
        pdf.close()
        resp = _upload(admin_client, pid, "chapter.pdf", buf.getvalue(), "正文")
        assert resp.status_code == 200, resp.text[:400]
        data = resp.json()
        assert len(data.get("chapters", [])) >= 1
